#!/usr/bin/env python3
"""
Generate LoopForgeLab Competitor Report as Word document.
"""

import io
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ──────────────────────────────────────────────
# Brand colours (LFL identity)
# ──────────────────────────────────────────────
DARK_BLUE   = RGBColor(0x0D, 0x2B, 0x4A)   # #0D2B4A  — headings / table header
MID_BLUE    = RGBColor(0x1A, 0x5F, 0x8A)   # #1A5F8A  — section headers
ACCENT      = RGBColor(0x00, 0xA8, 0xCC)   # #00A8CC  — accent / links
LIGHT_GREY  = RGBColor(0xF2, 0xF5, 0xF8)   # #F2F5F8  — table row bg
TEXT_DARK   = RGBColor(0x1A, 0x1A, 0x2E)   # #1A1A2E  — body text

PIE_COLOURS = ["#1A5F8A", "#00A8CC", "#F0A500", "#2ECC71"]

# ──────────────────────────────────────────────
# Helper: set paragraph shading
# ──────────────────────────────────────────────
def shade_cell(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color.lstrip('#'))
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'),   'single')
        tag.set(qn('w:sz'),    '4')
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), 'E0E8F0')
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def header_row(table, row_idx, hex_bg='0D2B4A', txt_color=RGBColor(0xFF,0xFF,0xFF)):
    row = table.rows[row_idx]
    for cell in row.cells:
        shade_cell(cell, hex_bg)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = txt_color
                run.font.bold = True

def alt_row(table, row_idx, hex_bg='F2F5F8'):
    row = table.rows[row_idx]
    for cell in row.cells:
        shade_cell(cell, hex_bg)

# ──────────────────────────────────────────────
# Helper: add formatted paragraph
# ──────────────────────────────────────────────
def add_para(doc, text, style='Normal', bold=False, size=None,
             color=None, align=None, space_before=None, space_after=None):
    p = doc.add_paragraph(style=style)
    if align:
        p.alignment = align
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p

def add_bullet(doc, text, level=0, size=10):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = TEXT_DARK
    return p

# ──────────────────────────────────────────────
# Pie chart generator → returns PNG bytes
# ──────────────────────────────────────────────
def make_pie_chart():
    labels  = ['Engineering\n270 (72.8%)',
               'Data & Req. Mgmt\n40 (10.8%)',
               'Sales & Quotes\n35 (9.4%)',
               'Sustainability\n26 (7.0%)']
    sizes   = [270, 40, 35, 26]
    explode = [0.03, 0.03, 0.08, 0.03]   # pop out Sales & Quotes

    fig, ax = plt.subplots(figsize=(7, 5), dpi=150)
    wedges, texts, autotexts = ax.pie(
        sizes, labels=labels, colors=PIE_COLOURS,
        explode=explode, autopct='%1.1f%%',
        startangle=140, pctdistance=0.75,
        textprops={'fontsize': 9, 'color': '#1A1A2E'},
        wedgeprops={'linewidth': 1.5, 'edgecolor': 'white'}
    )
    for at in autotexts:
        at.set_fontsize(8)
        at.set_color('white')
        at.set_fontweight('bold')

    ax.set_title('Marktanteile nach Unternehmenskategorie\n(Datenbasis: CI v1.7 | 371 Unternehmen)',
                 fontsize=11, fontweight='bold', color='#0D2B4A', pad=16)
    plt.tight_layout()

    buf = io.BytesIO()
    plt.savefig(buf, format='png', bbox_inches='tight')
    plt.close()
    buf.seek(0)
    return buf

# ──────────────────────────────────────────────
# MAIN
# ──────────────────────────────────────────────
def build_report():
    doc = Document()

    # ── Page margins ─────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Default font ─────────────────────────
    style = doc.styles['Normal']
    style.font.name  = 'Calibri'
    style.font.size  = Pt(10)
    style.font.color.rgb = TEXT_DARK

    # ══════════════════════════════════════════
    # TITLE PAGE
    # ══════════════════════════════════════════
    doc.add_paragraph()
    add_para(doc, 'LoopForgeLab', bold=True, size=28,
             color=DARK_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=30)
    add_para(doc, 'Competitive Intelligence Report', bold=True, size=20,
             color=MID_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4)
    add_para(doc, 'Markt- und Wettbewerbsanalyse · Fokus: Sales & Quotes', size=13,
             color=ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6)
    add_para(doc, ' ', size=10)
    add_para(doc, 'Datenbasis: Competitive Intelligence DB v1.7  |  Stand: März 2026',
             size=9, color=RGBColor(0x80,0x80,0x90), align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, 'Erstellt für: LoopForgeLab GbR · Berlin',
             size=9, color=RGBColor(0x80,0x80,0x90), align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # ══════════════════════════════════════════
    # 1 — EXECUTIVE SUMMARY / MARKTÜBERSICHT
    # ══════════════════════════════════════════
    doc.add_heading('1  Executive Summary & Marktübersicht', level=1)
    doc.paragraphs[-1].runs[0].font.color.rgb = DARK_BLUE

    add_para(doc, (
        'Dieser Competitive Intelligence Report analysiert das Wettbewerbsumfeld von LoopForgeLab '
        '— einem Berliner Deep-Tech-Start-up, das eine KI-gestützte "Product Intelligence Engine" '
        'für den mechanischen Produktentwicklungsprozess aufbaut. Die Datenbasis umfasst Version 1.7 '
        'der internen Competitive-Intelligence-Datenbank, erhoben und strukturiert auf Basis von '
        '44 qualitativen Experteninterviews sowie einer systematischen Marktrecherche.'
    ), space_after=6)

    # ── Summary KPI table ────────────────────
    doc.add_heading('Datenbasis auf einen Blick', level=2)
    doc.paragraphs[-1].runs[0].font.color.rgb = MID_BLUE

    kpi_tbl = doc.add_table(rows=2, cols=4)
    kpi_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    kpi_tbl.style = 'Table Grid'
    kpi_headers = ['Unique Companies', 'Produkte / Lösungen', 'Große Unternehmen', 'Start-ups']
    kpi_values  = ['371', '541', '76  (20 %)', '295  (80 %)']
    for i, (h, v) in enumerate(zip(kpi_headers, kpi_values)):
        hcell = kpi_tbl.cell(0, i)
        vcell = kpi_tbl.cell(1, i)
        hcell.text = h
        vcell.text = v
        shade_cell(hcell, '0D2B4A')
        shade_cell(vcell, 'E8F4F8')
        for para in hcell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9)
        for para in vcell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = DARK_BLUE
                run.font.size = Pt(14)
    doc.add_paragraph()

    # ── Category overview table ───────────────
    doc.add_heading('Marktaufteilung nach Kategorien', level=2)
    doc.paragraphs[-1].runs[0].font.color.rgb = MID_BLUE

    cat_data = [
        ('Engineering',                '270', '72.8 %', '380', '70.2 %', '54', '216', '80 %',
         'Factory Intelligence & IoT, BIM & Construction,\nEngineering Copilots & AI Assistants, Generative Design'),
        ('Data & Requirements Mgmt',   '40',  '10.8 %', '65',  '12.0 %', '7',  '33',  '82 %',
         'MBSE & Systems Engineering, PLM & PDM Platforms,\nEngineering Collaboration & Data, CAD & Modeling'),
        ('Sales & Quotes',             '35',  '9.4 %',  '39',  '7.2 %',  '13', '22',  '63 %',
         'Supply Chain & Procurement, CPQ, RfP/RfQ-Mgmt,\nE-Sourcing, Supply Chain Planning'),
        ('Sustainability',             '26',  '7.0 %',  '57',  '10.5 %', '2',  '24',  '92 %',
         'LCA Software & Platforms, Carbon & Eco-Design,\nCircular Economy, Regulatory & Compliance (DPP)'),
        ('GESAMT',                     '371', '100 %',  '541', '100 %',  '76', '295', '80 %', ''),
    ]
    col_headers = ['Kategorie', '# Co.', 'Anteil', '# Prod.', 'Anteil',
                   '# Groß', '# Start-up', 'Startup-%', 'Schwerpunkte']
    col_widths  = [Cm(3.5), Cm(1.2), Cm(1.3), Cm(1.3), Cm(1.3),
                   Cm(1.2), Cm(1.5), Cm(1.5), Cm(5.5)]

    cat_tbl = doc.add_table(rows=1 + len(cat_data), cols=len(col_headers))
    cat_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cat_tbl.style = 'Table Grid'
    for j, (hdr, w) in enumerate(zip(col_headers, col_widths)):
        cell = cat_tbl.cell(0, j)
        cell.width = w
        cell.text  = hdr
        shade_cell(cell, '0D2B4A')
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold  = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size  = Pt(8)

    row_colors = ['FFFFFF', 'F2F5F8', 'FFFFFF', 'F2F5F8', 'E8F4F8']
    for i, row_data in enumerate(cat_data):
        for j, val in enumerate(row_data):
            cell = cat_tbl.cell(i + 1, j)
            cell.text = val
            shade_cell(cell, row_colors[i])
            for para in cell.paragraphs:
                if j in (0,):
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(8)
                    if i == len(cat_data) - 1:  # totals row
                        run.font.bold = True
                        run.font.color.rgb = DARK_BLUE

    doc.add_paragraph()

    # ── Pie chart ────────────────────────────
    doc.add_heading('Verteilung der Unternehmen nach Kategorie', level=2)
    doc.paragraphs[-1].runs[0].font.color.rgb = MID_BLUE

    pie_buf = make_pie_chart()
    pic_para = doc.add_paragraph()
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pic_para.add_run()
    run.add_picture(pie_buf, width=Cm(13))
    doc.add_paragraph()

    add_para(doc, (
        'Die Engineering-Kategorie dominiert mit 72,8 % klar den Markt — ein Spiegelbild der breiten Digitalisierungswelle '
        'in CAD, CAE, IoT und KI-gestützten Entwicklungswerkzeugen. Sales & Quotes (9,4 %) und Sustainability (7,0 %) sind '
        'zahlenmäßig kleiner, weisen aber eine hohe Startup-Dichte auf, was auf intensive Innovationsdynamik und '
        'ungelöste Kernprobleme in diesen Bereichen hindeutet.'
    ), space_after=6)

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 2 — MARKTSEGMENTE: KURZPROFIL
    # ══════════════════════════════════════════
    doc.add_heading('2  Marktsegmente: Kurzprofil der vier Kategorien', level=1)
    doc.paragraphs[-1].runs[0].font.color.rgb = DARK_BLUE

    segments = [
        {
            'title': '2.1  Engineering  (270 Unternehmen | 380 Produkte)',
            'color': '1A5F8A',
            'body': (
                'Die Engineering-Kategorie ist mit Abstand die größte und breiteste Kategorie. '
                'Sie umfasst alle Werkzeuge, die direkt in den Produktentwicklungs- und '
                'Fertigungsprozess eingreifen: von klassischer CAD/CAE/CAM-Software über KI-gestützte '
                'Design-Copiloten bis hin zu Industrial-IoT-Plattformen, BIM-Lösungen für das '
                'Bauwesen und Factory-Intelligence-Systemen. '
                '80 % der 270 Unternehmen sind Startups — ein Zeichen für eine hochaktive '
                'Innovationslandschaft. Die etablierten Riesen (Autodesk, Siemens, Ansys, Dassault) '
                'konkurrieren mit einer Vielzahl agiler Newcomer, die spezifische Nischen in '
                'generativem Design, digitalen Zwillingen und KI-Assistenz besetzen.'
            ),
            'bullets': [
                'Sub-Kategorien: Factory Intelligence & IoT · BIM & Construction · Engineering Copilots & AI Assistants · Generative & Computational Design',
                'Relevanz für LFL: Direkte Überschneidung mit dem Design-Copilot-Ansatz von LoopForgeLab; Wettbewerber wie Autodesk Fusion und Siemens NX rüsten ihre Plattformen mit KI auf',
            ]
        },
        {
            'title': '2.2  Data & Requirements Management  (40 Unternehmen | 65 Produkte)',
            'color': '00A8CC',
            'body': (
                'Diese Kategorie adressiert das Kernproblem der fragmentierten Datenwelt in der '
                'Produktentwicklung: PLM-Plattformen, MBSE-Tools, Engineering-Kollaborations­lösungen '
                'und Anforderungsmanagement-Software. Mit nur 7 großen Unternehmen bei 40 insgesamt '
                '(82 % Startups) ist auch hier die Dynamik hoch. '
                'Etablierte Akteure wie PTC (Windchill), Dassault (ENOVIA) und SAP PLM dominieren den '
                'Enterprise-Markt, während Startups neue Ansätze für digitale Thread-Konzepte, '
                'kollaborative Anforderungsverwaltung und KI-gestützte Datenaggregation entwickeln.'
            ),
            'bullets': [
                'Sub-Kategorien: MBSE & Systems Engineering · PLM & PDM Platforms · Engineering Collaboration & Data · CAD & Modeling',
                'Relevanz für LFL: Der "Data Desert"-Befund aus dem Whitepaper (fragmentierte PLM/ERP-Daten) ist die zentrale Lücke, die LFL mit seiner Product Intelligence Engine schließen will',
            ]
        },
        {
            'title': '2.3  Sales & Quotes  (35 Unternehmen | 39 Produkte)',
            'color': 'F0A500',
            'body': (
                'Sales & Quotes ist die Kategorie mit dem direktesten Bezug zur Kernlösung von '
                'LoopForgeLab. Sie umfasst Lösungen rund um die Angebotsprozesse, den Einkauf und '
                'die Lieferkette: RfP/Bid-Response-Management, E-Sourcing & RfQ-Plattformen, '
                'CPQ-Systeme für komplexe Produkte und Supply-Chain-Planungs­lösungen. '
                'Mit 63 % Startup-Anteil ist diese Kategorie etwas ausgewogener zwischen Großen '
                'und Kleinen. Trotz vergleichsweise weniger Unternehmen ist die Problemlösungstiefe '
                'hoch — und der Schmerz der Kunden messbar groß.'
            ),
            'bullets': [
                'Sub-Kategorien: A — RfP/Bid Response Management · B — E-Sourcing/RfQ · C — CPQ (Configure Price Quote) · D — Supply Chain Planning',
                'Relevanz für LFL: LFLs RFQ-Modul (in Entwicklung ab Q2 2026) und die Angebots-Kostenanalyse sind direkte Antworten auf die Probleme in Sub-Cluster B und C',
                'Besonderes Merkmal: 39 Produkte bei nur 35 Unternehmen — viele Anbieter haben mehrere Produktlinien',
            ]
        },
        {
            'title': '2.4  Sustainability  (26 Unternehmen | 57 Produkte)',
            'color': '2ECC71',
            'body': (
                'Die kleinste Kategorie nach Unternehmenszahl, aber mit der höchsten Produkt­dichte '
                'pro Unternehmen (Ø 2,2 Produkte/Unternehmen). Sustainability umfasst LCA-Software, '
                'CO₂-Kalkulations­tools, Circular-Economy-Plattformen und Compliance-Lösungen '
                'für CSRD/DPP-Anforderungen. Mit 92 % Startup-Anteil ist dies der innovativste '
                'und am stärksten wachsende Bereich — getrieben durch den regulatorischen Druck der EU. '
                'Lediglich 2 etablierte Großunternehmen (CarbonBright, Sphera) haben hier bisher '
                'Fuss gefasst.'
            ),
            'bullets': [
                'Sub-Kategorien: LCA Software & Platforms · Carbon & Eco-Design · Circular Economy & R-Strategies · Regulatory & Compliance (DPP)',
                'Relevanz für LFL: LFLs "Carbon Case" (10,5 Mio. tCO₂e Einsparung bis Jahr 10) und die Echtzeit-CO₂-Feedback-Funktion positionieren LFL als Brücke zwischen Engineering und Sustainability',
                'Regulatorischer Treiber: CSRD-Pflicht ab 2025 für ~50.000 EU-Unternehmen erzeugt enormen Handlungsdruck',
            ]
        },
    ]

    for seg in segments:
        doc.add_heading(seg['title'], level=2)
        doc.paragraphs[-1].runs[0].font.color.rgb = RGBColor(
            int(seg['color'][0:2], 16), int(seg['color'][2:4], 16), int(seg['color'][4:6], 16)
        )
        add_para(doc, seg['body'], space_after=4)
        for bullet in seg['bullets']:
            add_bullet(doc, bullet)
        doc.add_paragraph()

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 3 — DEEP DIVE: SALES & QUOTES
    # ══════════════════════════════════════════
    doc.add_heading('3  Deep Dive: Sales & Quotes', level=1)
    doc.paragraphs[-1].runs[0].font.color.rgb = DARK_BLUE

    add_para(doc, (
        'Die folgende Tiefenanalyse der Kategorie Sales & Quotes kombiniert die strukturierten '
        'Marktdaten aus der CI-Datenbank v1.7 mit den qualitativen Erkenntnissen aus dem '
        'LoopForgeLab Whitepaper ("Optimized for Performance, Approved for Price", Feb. 2026) '
        'und dem Pitch Deck (März 2026). Ziel ist es, die Kundenproblemfelder messbar zu machen '
        'und zu zeigen, wo und wie LoopForgeLabs Lösung in den Markt greift.'
    ), space_after=8)

    # ── 3.1 Marktstruktur ─────────────────────
    doc.add_heading('3.1  Marktstruktur: Die vier Sub-Cluster', level=2)
    doc.paragraphs[-1].runs[0].font.color.rgb = MID_BLUE

    add_para(doc, (
        'Die Kategorie Sales & Quotes gliedert sich in vier funktionale Sub-Cluster, '
        'die unterschiedliche Kunden, Technologien und Probleme adressieren:'
    ), space_after=4)

    cluster_data = [
        ('A', 'RfP / Bid Response Management (KI-gestützt)', '8',
         'Generative AI (GPT-4/Claude), RAG, AI Agents, CRM-Integration',
         'B2B-Technologie, Engineering-DL, SaaS, Verteidigung, Beratung'),
        ('B', 'E-Sourcing / RfQ-Management (Einkauf)', '8',
         'Autonomous AI Sourcing Bots, Reverse Auctions, ERP-Integration (SAP/Oracle), Spend Analytics',
         'Enterprise/Mid-Market Einkauf: Fertigung, Pharma, Automotive, Öffentl. Hand'),
        ('C', 'CPQ – Configure Price Quote', '3',
         'Constraint-based AI Configuration, CAD-Integration, ERP/CRM-Kopplung, CO₂-Kalkulation',
         'Hersteller von ETO-Produkten: Industriemaschinen, Automotive-Zulieferer, High-Tech'),
        ('D', 'Supply Chain Planning & Optimierung', '14',
         'AI/ML-Demand Forecasting, Digital Twins, Concurrent Planning, Graph Neural Networks, Math. Optimization',
         'Automotive, Consumer Goods, Pharma, Aerospace, Retail, High-Tech'),
    ]

    cl_tbl = doc.add_table(rows=1 + len(cluster_data), cols=5)
    cl_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cl_tbl.style = 'Table Grid'
    cl_headers = ['ID', 'Sub-Cluster', '# Co.', 'Schlüssel-Technologien', 'Haupt-Zielbranchen']
    cl_widths  = [Cm(0.8), Cm(4.5), Cm(1.0), Cm(5.5), Cm(5.5)]
    for j, (h, w) in enumerate(zip(cl_headers, cl_widths)):
        cell = cl_tbl.cell(0, j)
        cell.width = w
        cell.text  = h
        shade_cell(cell, '1A5F8A')
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold  = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size  = Pt(8)

    cl_colors = ['FFFFFF', 'F2F5F8', 'FFFFFF', 'F2F5F8']
    for i, row in enumerate(cluster_data):
        for j, val in enumerate(row):
            cell = cl_tbl.cell(i + 1, j)
            cell.text = val
            shade_cell(cell, cl_colors[i])
            for para in cell.paragraphs:
                if j == 0:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(8)
                    if j == 0:
                        run.font.bold = True
                        run.font.color.rgb = MID_BLUE

    doc.add_paragraph()

    # ── 3.2 Kundenproblemfelder ──────────────
    doc.add_heading('3.2  Kundenproblemfelder — und der LFL-Kontext', level=2)
    doc.paragraphs[-1].runs[0].font.color.rgb = MID_BLUE

    add_para(doc, (
        'Das Whitepaper von LoopForgeLab dokumentiert vier strukturelle Barrieren für '
        'profitable und nachhaltige Produktentwicklung. Diese Barrieren spiegeln sich '
        'unmittelbar in den gemessenen Kundenproblemen der Sales-&-Quotes-Kategorie wider:'
    ), space_after=4)

    # Intro framing boxes
    framing = [
        ('Der "Economic Gatekeeper"',
         '"Der Preis entscheidet über alles und jeden" — so das Credo aus 44 LFL-Experteninterviews. '
         'Dieser Befund schlägt sich direkt in der CPQ- und RfQ-Welt nieder: Engineers verbringen '
         '30–40 % ihrer Zeit mit Angebotserstellung statt mit Design — weil Konfiguration und '
         'Preisbildung manuell, fehleranfällig und vom CAD-Modell entkoppelt sind.'),
        ('Die "Data Desert"',
         'Kritische Produktdaten — CO₂-Fußabdrücke, Lieferantenpreise, Fertigungsbeschränkungen — '
         'sind über PLM, ERP, Excel und lokale Dokumente verstreut. Das Ergebnis: RfP-Teams '
         'schreiben 30–40 % aller Antworten von Grund auf neu, obwohl die Informationen '
         'irgendwo im Unternehmen existieren. Der Loopio-Benchmark misst 17+ Stunden pro RfP.'),
        ('Die "Redesign Trap"',
         'Weil Angebot und CAD-Konstruktion vollständig entkoppelt sind, entstehen '
         '50–70 % aller Engineering-Änderungen durch Inkonsistenz zwischen Angebot und CAD '
         '(DriveWorks-Studien). Jede späte Änderung multipliziert die Kosten — typisch '
         '5–15 % des Auftragswertes für Fehlerkorrektur nach Auftragsvergabe.'),
        ('Der "Tribal Knowledge"-Verlust',
         'Wenn erfahrene Ingenieure in Rente gehen, verschwindet das implizite Wissen, '
         'warum bestimmte Konfigurationen funktionieren oder warum bestimmte Lieferanten '
         'bevorzugt werden. CPQ-Systeme können dieses Wissen nur dann skalieren, '
         'wenn es explizit kodifiziert ist — genau das ist heute nicht der Fall.'),
    ]

    for title, body in framing:
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(3)
        p.paragraph_format.space_before = Pt(6)
        r = p.add_run(f'■  {title}')
        r.bold = True
        r.font.color.rgb = DARK_BLUE
        r.font.size = Pt(10)
        add_para(doc, body, space_after=2)

    doc.add_paragraph()

    # ── Problem table ─────────────────────────
    doc.add_heading('Konkrete Kundenproblemfelder nach Sub-Cluster', level=3)
    doc.paragraphs[-1].runs[0].font.color.rgb = MID_BLUE

    prob_data = [
        ('A – RfP/Bid Response', 'Zeitaufwand',       '17+ h/RfP Beantwortung (Loopio-Benchmark)',            '68 h/Monat = ~1,7 FTE nur für RfP-Arbeit',                '★★★★★'),
        ('A – RfP/Bid Response', 'Wissenssilos',      '30–40 % aller Antworten neu verfasst, Wissen existiert','Massive Doppelarbeit; kein institutionelles Gedächtnis',   '★★★★★'),
        ('A – RfP/Bid Response', 'Bid/No-Bid',        'Entscheidung bauchgefühl-basiert, kein Score',         '80–85 % der Aufwände für verlorene Bids',                 '★★★★☆'),
        ('A – RfP/Bid Response', 'Win-Theme-Lücke',   'Kein Zugriff auf CRM/Gesprächsdaten für Positionierung','Generische Angebote → niedrige Win-Rate',                '★★★★☆'),
        ('B – E-Sourcing/RfQ',   'Excel/E-Mail-Chaos','RfX-Prozesse laufen über Excel ohne Versionierung',     '3–5 Tage Durchlaufzeit; häufige Fehler',                  '★★★★★'),
        ('B – E-Sourcing/RfQ',   'CSRD/Scope-3',      '~50.000 EU-Unternehmen müssen Scope-3 berichten',       'Manuelle Lieferantenbefragung für hunderte Tier-1-Lieferanten','★★★★★'),
        ('B – E-Sourcing/RfQ',   'KMU-Zugangshürde',  'Enterprise-Suites kosten €100K+/Jahr',                 '99,8 % europ. Fertigungsunternehmen sind KMU ohne digitales Sourcing','★★★★★'),
        ('C – CPQ',              'Ingenieur-Bottleneck','Jedes Angebot erfordert Ingenieurbeteiligung',         'ETO-Angebote: 2–4 Wochen; Ingenieure 30–40 % Zeit für Angebote','★★★★★'),
        ('C – CPQ',              'Kein CAD-Angebots-Link','Angebot und CAD vollständig entkoppelt',            '50–70 % Engineering Changes durch Inkonsistenz',          '★★★★☆'),
        ('C – CPQ',              'ETO ohne Lösung',    'Keine Standard-CPQ für echte ETO-Produkte',            '~60 % aller EU-Maschinenbauaufträge sind kundenspezifisch','★★★★★'),
        ('D – SC Planning',      'Forecast-Fehler',   'Excel + Legacy-APS, kein Echtzeit-Update',              'Ø 30–40 % MAPE Forecast-Fehler; hohe Bestandskosten',     '★★★★★'),
        ('D – SC Planning',      'Multi-Tier-Blindheit','Nur Tier-1-Lieferanten sichtbar',                     '77 % der Disruptions entstehen bei Tier-2+ Lieferanten',  '★★★★☆'),
    ]
    prob_headers = ['Sub-Cluster', 'Problem', 'Beschreibung', 'Messbare Auswirkung', 'Priorität']
    prob_widths  = [Cm(3.0), Cm(2.8), Cm(5.5), Cm(5.0), Cm(1.5)]

    prob_tbl = doc.add_table(rows=1 + len(prob_data), cols=5)
    prob_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    prob_tbl.style = 'Table Grid'

    for j, (h, w) in enumerate(zip(prob_headers, prob_widths)):
        cell = prob_tbl.cell(0, j)
        cell.width = w
        cell.text  = h
        shade_cell(cell, '0D2B4A')
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold  = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size  = Pt(8)

    prev_cluster = None
    for i, row in enumerate(prob_data):
        bg = 'FFFFFF' if i % 2 == 0 else 'F2F5F8'
        for j, val in enumerate(row):
            cell = prob_tbl.cell(i + 1, j)
            cell.text = val
            shade_cell(cell, bg)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8)
                    if j == 0:
                        run.font.bold = True
                        run.font.color.rgb = MID_BLUE

    doc.add_paragraph()

    # ── 3.3 LFL-Bezug konkret ────────────────
    doc.add_heading('3.3  Relevanz für LoopForgeLab — der strategische Fit', level=2)
    doc.paragraphs[-1].runs[0].font.color.rgb = MID_BLUE

    add_para(doc, (
        'LoopForgeLab positioniert sich mit seiner "Product Intelligence Engine" als wirtschaftliche '
        'Integrationsschicht über dem Engineering-Ökosystem. Das Pitch Deck identifiziert vier '
        'direkte Wettbewerbsfelder: Engineering, Data & Requirements Management, Sustainability und '
        'Sales & Quotes. Aus den gemessenen Kundenproblemen ergibt sich ein klarer strategischer Fit:'
    ), space_after=6)

    fit_data = [
        ('RFQ-Modul (ab Q2 2026)',
         'Sub-Cluster B & C',
         'LFLs erstes Kernmodul automatisiert die Angebotskalkulation für mechanische Produkte. '
         'Es löst den "Ingenieur-Bottleneck" (30–40 % Ingenieurszeit für Angebote) und schließt '
         'die Lücke zwischen CAD-Modell und Angebotskonfiguration — das Kernproblem von CPQ für ETO.'),
        ('Manufacturing Cost Intelligence',
         'Sub-Cluster B, C',
         'Echtzeit-Fertigungskosten direkt aus Engineering-Daten — ohne Excel-Hell. '
         'Adressiert den "Economic Gatekeeper"-Befund: Entscheidungen können früher '
         'im Design-Prozess auf Kostenbasis getroffen werden.'),
        ('Compliance & Risk',
         'Sub-Cluster B, D',
         'Integration von CSRD/Scope-3-Anforderungen in den Design- und Beschaffungsprozess. '
         'Adressiert den wachsenden regulatorischen Druck, der ab 2025 für ~50.000 EU-Unternehmen gilt.'),
        ('Life Cycle Cost & Operator Insights',
         'Sub-Cluster C, D',
         'Total Cost of Ownership, Reparierbarkeit und Kreislaufwirtschafts-Potenziale '
         'als native Designparameter — Sustainability als Innovationstreiber statt Compliance-Bürde.'),
    ]

    for title, cluster, body in fit_data:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(2)
        r1 = p.add_run(f'▶  {title}')
        r1.bold = True
        r1.font.color.rgb = DARK_BLUE
        r1.font.size = Pt(10)
        r2 = p.add_run(f'   [{cluster}]')
        r2.font.color.rgb = ACCENT
        r2.font.size = Pt(9)
        add_para(doc, body, space_after=2)

    doc.add_paragraph()

    # Market size context
    add_para(doc,
        'Marktgröße (Pitch Deck):  TAM € 16,6 Mrd. · SAM € 5,6 Mrd. · SOM € 98 Mio.  '
        '(415.000 Maschinenbauunternehmen, CAGR Engineeringsoftware 15 %)',
        bold=True, color=DARK_BLUE, space_after=2)
    add_para(doc,
        'Primäre Zielmärkte: Verpackungsmaschinen-Produzenten · Material Handling Machinery · '
        'Engineering Services — OEM in DE/AT/CH/PL/CZ/NL/BE mit ≥ 30 Mitarbeitern',
        space_after=8)

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 4 — LÖSUNGSANBIETER & TECHNOLOGIEN
    # ══════════════════════════════════════════
    doc.add_heading('4  Lösungsanbieter & Technologien in Sales & Quotes', level=1)
    doc.paragraphs[-1].runs[0].font.color.rgb = DARK_BLUE

    # ── 4.1 Anbieterprofile ──────────────────
    doc.add_heading('4.1  Anbieterprofile nach Sub-Cluster', level=2)
    doc.paragraphs[-1].runs[0].font.color.rgb = MID_BLUE

    # Cluster A
    doc.add_heading('Sub-Cluster A — RfP / Bid Response Management', level=3)
    doc.paragraphs[-1].runs[0].font.color.rgb = RGBColor(0x1A, 0x5F, 0x8A)

    add_para(doc, (
        'Dieser Sub-Cluster konzentriert sich auf die KI-gestützte Automatisierung von '
        'Antworten auf Ausschreibungen (RfP, RfI, DDQ). Der Markt ist 2023–2025 explodiert '
        'und wird von einer Welle KI-nativer Startups aufgemischt, die Legacy-Anbieter wie '
        'Loopio herausfordern.'
    ), space_after=4)

    a_data = [
        ('Loopio', 'Groß', 'RfP-Marktführer mit 1.500+ Kunden. KI auf 10+ Jahren Proposal-Daten trainiert. Stärken: UX, Support (9,7/10 G2). Schwäche: Library-basierte KI, keine Live-Verbindungen, langsamer als neue AI-Native-Tools.'),
        ('Responsive.io', 'Groß', '2.000+ Kunden, 100+ Integrationen (Salesforce, Slack, MS Teams), 50+ Sprachen. Stärken: Integrations-Ökosystem, Analytics. Schwäche: Steile Lernkurve, komplexes Pricing.'),
        ('Arphie', 'Startup', 'KI-nativ mit Live-KB-Verbindungen (Google Drive, SharePoint, Confluence). 80 %+ Zeitersparnis vs. Loopio-Benchmark, 84 % AI Acceptance Rate. Flat-Rate-Pricing (unlimitierte User). Zero Data Retention als Trust-Feature.'),
        ('Altura', 'Startup', 'Autonomer "Bid Companion" KI-Agent, der RfP-Aufgaben selbstständig erledigt und Compliance-Risiken proaktiv kennzeichnet. Europäischer Fokus, Echtzeit Bid/No-Bid Intelligence.'),
        ('AutogenAI', 'Startup', '$36,1M Umsatz (2025), 200+ Enterprise-Kunden. Generiert komplette Bid-Dokumente für komplexe Government- und Engineering-RfPs. Speziell für narrative Heavy-Bids, weniger für strukturierte RfQ-Formulare.'),
        ('Inventive AI', 'Startup', 'Identifiziert "Win Themes" aus CRM + Slack + Sales Calls für strategisch überlegene RfP-Antworten. Deep CRM/Salesforce-Integration. $1,7M ARR bootstrapped.'),
        ('DeepRFP', 'Startup', 'KI-Agent-Plattform mit 28-Sprachen-Support und dedizierten Compliance-Automatisierungs-Agenten. Umfassendster multilingualer Ansatz im Cluster.'),
        ('AutoRFP.ai', 'Startup', 'Browser-basierte RfP-Automation via Chrome Extension — arbeitet direkt in SAP Ariba und Government-Portalen. Unique Workflow für beschaffungsgetriebene Organisationen.'),
    ]

    a_tbl = doc.add_table(rows=1 + len(a_data), cols=3)
    a_tbl.style = 'Table Grid'
    a_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, (h, w) in enumerate(zip(['Unternehmen', 'Typ', 'Lösung / USP'],
                                    [Cm(3.0), Cm(1.5), Cm(13.0)])):
        cell = a_tbl.cell(0, j)
        cell.width = w
        cell.text  = h
        shade_cell(cell, '1A5F8A')
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold  = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size  = Pt(8)

    for i, (name, typ, desc) in enumerate(a_data):
        bg = 'FFFFFF' if i % 2 == 0 else 'F2F5F8'
        for j, val in enumerate([name, typ, desc]):
            cell = a_tbl.cell(i + 1, j)
            cell.text = val
            shade_cell(cell, bg)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8)
                    if j == 0:
                        run.font.bold = True
                        run.font.color.rgb = DARK_BLUE
                    elif j == 1:
                        run.font.color.rgb = (MID_BLUE if val == 'Startup'
                                              else RGBColor(0xC0,0x39,0x2B))

    doc.add_paragraph()

    # Cluster B
    doc.add_heading('Sub-Cluster B — E-Sourcing / RfQ-Management (Einkauf)', level=3)
    doc.paragraphs[-1].runs[0].font.color.rgb = RGBColor(0x1A, 0x5F, 0x8A)

    add_para(doc, (
        'E-Sourcing-Plattformen digitalisieren den Einkaufs- und Beschaffungsprozess. '
        'Der Markt ist zweigeteilt: Enterprise-Suites für Großkonzerne (Coupa, Ivalua, Jaggaer) '
        'und erschwingliche Alternativen für den Mittelstand. Das drängendste Problem bleibt '
        'die Unerschwinglichkeit großer Suites für KMU.'
    ), space_after=4)

    b_data = [
        ('Coupa', 'Groß', '$619,4M Umsatz (FY2025), 3.500 Mitarbeiter. Community Intelligence aus $6T+ Transaktionen für Benchmarking und Risikodetektion. End-to-End BSM. Schwäche: Premium-Pricing, KMU-ungeeignet.'),
        ('Ivalua', 'Groß', 'Globaler S2P-Marktführer auf einheitlicher Plattform. Starke EU-Präsenz, CSRD-Reporting-Integration. Hohe Implementierungskosten.'),
        ('Jaggaer', 'Groß', 'Autonome Procurement-Agenten (JAI) für automatisierten Source-to-Pay. Starke Fertigungsbranchenausrichtung. Enterprise-only.'),
        ('Keelvar', 'Startup', 'Einziger Anbieter im Gartner Market Guide 2025 für SOWOHL Advanced Sourcing Optimization ALS AUCH Autonomous Sourcing. $90B+ Spend managed. Autonome AI-Bots führen komplette RfQ-Zyklen ohne manuelle Intervention durch.'),
        ('ProcurePort', 'Startup', 'Erschwingliche Full-Suite eProcurement seit 2011 (bootstrapped). Handhabt 1.000–15.000-Item-RfX-Events. Signifikant günstiger als Coupa/Ivalua/Jaggaer — ideal für KMU-Einstieg.'),
        ('QLM Sourcing', 'Startup', 'Custom eRfQ-Template-Engine: Pro-Produkt-Kategorie anpassbare RfQ-Formulare für Engineering-Einkauf. Starke Lieferanten-Kollaboration.'),
        ('Bonfire', 'Startup', '600+ Public-Sector-Procurement-Teams, $18,2M Umsatz. Digitale Scorecards, gewichtete Bewertung, What-If-Analyse. Schwäche: Primär öffentlicher Sektor.'),
        ('Vendorful', 'Startup', 'API-first E-Sourcing mit tiefer ERP-Integration, ersetzt Excel/SharePoint-RfX-Workflows. Für Beginner und Experten geeignet.'),
    ]

    b_tbl = doc.add_table(rows=1 + len(b_data), cols=3)
    b_tbl.style = 'Table Grid'
    b_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, (h, w) in enumerate(zip(['Unternehmen', 'Typ', 'Lösung / USP'],
                                    [Cm(3.0), Cm(1.5), Cm(13.0)])):
        cell = b_tbl.cell(0, j)
        cell.width = w
        cell.text  = h
        shade_cell(cell, '1A5F8A')
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold  = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size  = Pt(8)

    for i, (name, typ, desc) in enumerate(b_data):
        bg = 'FFFFFF' if i % 2 == 0 else 'F2F5F8'
        for j, val in enumerate([name, typ, desc]):
            cell = b_tbl.cell(i + 1, j)
            cell.text = val
            shade_cell(cell, bg)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8)
                    if j == 0:
                        run.font.bold = True
                        run.font.color.rgb = DARK_BLUE
                    elif j == 1:
                        run.font.color.rgb = (MID_BLUE if val == 'Startup'
                                              else RGBColor(0xC0, 0x39, 0x2B))

    doc.add_paragraph()

    # Cluster C
    doc.add_heading('Sub-Cluster C — CPQ: Configure Price Quote', level=3)
    doc.paragraphs[-1].runs[0].font.color.rgb = RGBColor(0x1A, 0x5F, 0x8A)

    add_para(doc, (
        'CPQ ist der kleinste Sub-Cluster (3 Unternehmen), aber mit dem stärksten strategischen '
        'Fit zu LoopForgeLab. Die Kernaufgabe: komplexe, kundenspezifische Produkte '
        '(Engineer-to-Order) so konfigurieren, dass ein korrektes Angebot in Minuten statt '
        'Wochen entsteht — inkl. CAD-Generierung. Rund 60 % aller Aufträge im europäischen '
        'Maschinenbau sind kundenspezifisch; eine skalierbare Standardlösung fehlt bisher.'
    ), space_after=4)

    c_data = [
        ('Tacton', 'Groß', '301 Mitarbeiter, Stockholm. Marktführer CPQ für komplexe Fertigungsprodukte. Einzige CPQ-Lösung mit integrierter CO₂-Fußabdruckberechnung und EPD-Generierung direkt in der Produktkonfiguration. Starke CAD- und ERP/CRM-Kopplung.'),
        ('DriveWorks', 'Startup', 'Sheffield, UK. Design Automation + CPQ nativ für SolidWorks. Regelbasierte CAD-Generierung direkt aus dem Angebot — einzigartiger Workflow: Angebot → CAD-Modell automatisch. Dominant im SolidWorks-Ökosystem.'),
        ('Elfsquad', 'Startup', 'Groningen, NL. CPQ/Konfigurations-Software speziell für Fertigungs- und Industrieunternehmen. Fokus auf visuelle Produktkonfiguration und nahtlose ERP-Integration für den DACH/Benelux-Mittelstand.'),
    ]

    c_tbl = doc.add_table(rows=1 + len(c_data), cols=3)
    c_tbl.style = 'Table Grid'
    c_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, (h, w) in enumerate(zip(['Unternehmen', 'Typ', 'Lösung / USP'],
                                    [Cm(3.0), Cm(1.5), Cm(13.0)])):
        cell = c_tbl.cell(0, j)
        cell.width = w
        cell.text  = h
        shade_cell(cell, '1A5F8A')
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold  = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size  = Pt(8)

    for i, (name, typ, desc) in enumerate(c_data):
        bg = 'FFFFFF' if i % 2 == 0 else 'F2F5F8'
        for j, val in enumerate([name, typ, desc]):
            cell = c_tbl.cell(i + 1, j)
            cell.text = val
            shade_cell(cell, bg)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8)
                    if j == 0:
                        run.font.bold = True
                        run.font.color.rgb = DARK_BLUE
                    elif j == 1:
                        run.font.color.rgb = (MID_BLUE if val == 'Startup'
                                              else RGBColor(0xC0, 0x39, 0x2B))

    doc.add_paragraph()

    # Cluster D
    doc.add_heading('Sub-Cluster D — Supply Chain Planning & Optimierung', level=3)
    doc.paragraphs[-1].runs[0].font.color.rgb = RGBColor(0x1A, 0x5F, 0x8A)

    add_para(doc, (
        'Mit 14 Unternehmen ist dies der größte Sub-Cluster. Supply-Chain-Planning-Lösungen '
        'adressieren Forecast-Fehler, Echtzeitreaktionsfähigkeit und Multi-Tier-Lieferantensichtbarkeit. '
        'Der Markt wird von wenigen großen Plattformanbietern dominiert, ergänzt durch spezialisierte '
        'AI-Startups für Risikomanagement und Optimierung.'
    ), space_after=4)

    d_data = [
        ('Kinaxis', 'Groß', 'Ottawa, CA. Concurrent Planning Platform "Maestro" — einzige AI-infused End-to-End SC-Orchestrierungsplattform. Echtzeit-Plananpassung statt sequenzieller Zyklen. Dominant in Automotive und High-Tech.'),
        ('Blue Yonder', 'Groß', 'Scottsdale, USA. AI/ML-powered Supply Chain Planning & Execution. Breite Branchenabdeckung. Starke Demand-Sensing-Kapazitäten.'),
        ('o9 Solutions', 'Startup', 'Dallas, TX. Proprietäres Enterprise Knowledge Graph (EKG) für End-to-End SC-Modellierung. Neurosymbolische AI Agents + GenAI/LLM Composite Agents für Cross-funktionale Planung.'),
        ('Resilinc', 'Startup', 'Milpitas, CA. Agentic AI auf 15+ Jahren SC-Daten. Autonome Risikoerkennung und -minderung. Spezialist für Multi-Tier-Sichtbarkeit und Disruption-Response.'),
        ('Scoutbee', 'Startup', 'Würzburg, DE. KI-gestützte Lieferantenentdeckung und -risikomanagement mit LLMs. Europäischer Player, starker Bezug zur deutschen Fertigungsindustrie.'),
        ('Cosmo Tech', 'Startup', 'Lyon, FR. Simulations-basierte Supply-Chain-Optimierung mit Digital Twins. Stärke: komplexe Szenarioanalysen für non-lineare SC-Probleme.'),
    ]

    d_tbl = doc.add_table(rows=1 + len(d_data), cols=3)
    d_tbl.style = 'Table Grid'
    d_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, (h, w) in enumerate(zip(['Unternehmen', 'Typ', 'Lösung / USP'],
                                    [Cm(3.0), Cm(1.5), Cm(13.0)])):
        cell = d_tbl.cell(0, j)
        cell.width = w
        cell.text  = h
        shade_cell(cell, '1A5F8A')
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold  = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size  = Pt(8)

    for i, (name, typ, desc) in enumerate(d_data):
        bg = 'FFFFFF' if i % 2 == 0 else 'F2F5F8'
        for j, val in enumerate([name, typ, desc]):
            cell = d_tbl.cell(i + 1, j)
            cell.text = val
            shade_cell(cell, bg)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8)
                    if j == 0:
                        run.font.bold = True
                        run.font.color.rgb = DARK_BLUE
                    elif j == 1:
                        run.font.color.rgb = (MID_BLUE if val == 'Startup'
                                              else RGBColor(0xC0, 0x39, 0x2B))

    doc.add_paragraph()
    doc.add_page_break()

    # ── 4.2 Technologie-Radar ────────────────
    doc.add_heading('4.2  Technologie-Radar: Schlüsseltechnologien in Sales & Quotes', level=2)
    doc.paragraphs[-1].runs[0].font.color.rgb = MID_BLUE

    add_para(doc, (
        'Der folgende Technologie-Radar zeigt die wichtigsten Technologien im Sales-&-Quotes-Markt, '
        'ihren Reifegrad und ihre Relevanz für LoopForgeLab:'
    ), space_after=4)

    tech_data = [
        ('Generative AI (LLM) für Proposal-Texte', '★★★★★ Produktiv', 'A',
         'Arphie, AutogenAI, Inventive AI, Loopio, DeepRFP, Altura',
         'Entscheidend — Qualität, Geschwindigkeit, Personalisierung',
         'Direkt relevant für LFLs RFQ-Textgenerierung'),
        ('RAG (Retrieval-Augmented Generation)', '★★★★★ Produktiv', 'A+B',
         'Arphie (live KB), Inventive AI, Responsive.io',
         'Basis für AI-native Tools — Wissensbasiszugriff',
         'Kern-Technologie für LFLs Data-Desert-Lösung'),
        ('Autonomous AI Agents', '★★★★☆ Aufsteigend', 'A+B',
         'Keelvar (Sourcing Bots), Altura (Bid Companion), Arphie',
         'Nächste Evolutionsstufe: von Assist zu Autonom',
         'Roadmap-relevant für LFLs Design Copilot'),
        ('Constraint-based Configuration Engine', '★★★★★ Produktiv', 'C',
         'Tacton, DriveWorks, Elfsquad',
         'Kern-IP für CPQ bei ETO — Differenzierung über Komplexitätstiefe',
         'Direkte Konkurrenz zum LFL-RFQ-Modul für ETO'),
        ('CAD-CPQ-Integration', '★★★★☆ Produktiv', 'C',
         'DriveWorks (SolidWorks-nativ)',
         'Einzigartiger Workflow: Angebot → CAD-Modell automatisch',
         'Kern-Feature des LFL-Differenzierungsansatzes'),
        ('ML für Instant Quoting', '★★★★★ Produktiv', 'B (Fertigung)',
         'Xometry, Protolabs',
         'ML auf Millionen Teilen trainiert — Preis in <60 Sek.',
         'Langfristiges Ziel für LFL im Fertigungsnetzwerk'),
        ('Concurrent Planning / Digital Twin SC', '★★★★☆ Aufsteigend', 'D',
         'Kinaxis (RapidResponse), Cosmo Tech, o9 Solutions',
         'Echtzeit-Plananpassung ohne sequentielle Zyklen',
         'Mittelbar relevant für LFLs Supply-Chain-Datenschicht'),
        ('ESG/CSRD-Daten in Sourcing-Workflow', '★★★☆☆ Wachsend', 'B+C',
         'Tacton (CO₂-Konfigurator), Jaggaer, Ivalua',
         'Pflicht-Feature für EU-Compliance ab 2025',
         'Kernbestandteil der LFL Sustainability-Positionierung'),
        ('Bid/No-Bid AI Scoring', '★★★☆☆ Aufsteigend', 'A',
         'Altura, Inventive AI, AutogenAI',
         'Ressourcenallokation optimieren: richtige Bids priorisieren',
         'Relevant für LFLs Technical Sales Zielgruppe'),
    ]

    tech_headers = ['Technologie', 'Reifegrad', 'Cluster', 'Beispiel-Anbieter',
                    'Wettbewerbs-Differenzierung', 'LFL-Relevanz']
    tech_widths  = [Cm(3.5), Cm(2.0), Cm(1.0), Cm(3.8), Cm(4.0), Cm(3.5)]

    tech_tbl = doc.add_table(rows=1 + len(tech_data), cols=6)
    tech_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tech_tbl.style = 'Table Grid'

    for j, (h, w) in enumerate(zip(tech_headers, tech_widths)):
        cell = tech_tbl.cell(0, j)
        cell.width = w
        cell.text  = h
        shade_cell(cell, '0D2B4A')
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold  = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size  = Pt(8)

    for i, row in enumerate(tech_data):
        bg = 'FFFFFF' if i % 2 == 0 else 'F2F5F8'
        for j, val in enumerate(row):
            cell = tech_tbl.cell(i + 1, j)
            cell.text = val
            shade_cell(cell, bg)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8)
                    if j == 5:  # LFL relevanz column
                        run.font.color.rgb = RGBColor(0x1A, 0x5F, 0x8A)
                        run.font.italic = True

    doc.add_paragraph()
    doc.add_page_break()

    # ══════════════════════════════════════════
    # 5 — POSITIONIERUNG & WHITESPACE
    # ══════════════════════════════════════════
    doc.add_heading('5  Strategische Positionierung & Whitespace für LoopForgeLab', level=1)
    doc.paragraphs[-1].runs[0].font.color.rgb = DARK_BLUE

    add_para(doc, (
        'Aus der Analyse der 35 Unternehmen und 39 Lösungen in der Sales-&-Quotes-Kategorie '
        'ergibt sich ein klares Bild über Marktlücken, die LoopForgeLab gezielt adressieren kann:'
    ), space_after=6)

    whitespace = [
        ('Engineering-nativer RFQ-Prozess fehlt',
         'Alle bestehenden CPQ- und E-Sourcing-Lösungen setzen nach dem Design-Prozess an. '
         'Keine Lösung verankert die Angebotskalkulation nativ im CAD-basierten Engineering-Workflow. '
         'Tacton und DriveWorks kommen am nächsten, sind aber nicht als "Engineering-First"-Werkzeug positioniert. '
         'LFL schließt diese Lücke mit seiner Product Intelligence Engine — Kosten, Compliance und CO₂ '
         'werden zu nativen Designparametern, nicht zu nachgelagerten Berechnungen.'),
        ('KMU-taugliche ETO-CPQ fehlt',
         'Die drei CPQ-Anbieter (Tacton, DriveWorks, Elfsquad) sind entweder Enterprise-only oder '
         'SolidWorks-exklusiv. Für den europäischen Mittelstand (Maschinenbau, Material Handling) '
         'gibt es keine erschwingliche, CAD-agnostische CPQ-Lösung für echte ETO-Produkte. '
         'LFLs Zielsegment (OEM >30 MA in DACH/Benelux) ist genau diese Lücke.'),
        ('CO₂ + Kosten in einem Tool',
         'Nur Tacton hat ansatzweise einen CO₂-Kalkulatoren in CPQ integriert. '
         'Die Verbindung von Real-time-Kostenberechnung, CO₂-Feedback und Angebotsgenerierung '
         'in einer für Engineers konzipierten Oberfläche existiert nicht. '
         'LFLs "Carbon Case" (bis zu 10,5 Mio. tCO₂e Einsparung im Jahr 10 bei 14 % Marktanteil) '
         'positioniert diese Integration als strategischen Differenziator.'),
        ('Tribal Knowledge als Wettbewerbsvorteil',
         'Kein Anbieter in der Sales-&-Quotes-Kategorie adressiert explizit die Kodifizierung '
         'von institutionellem Ingenieurs-Wissen. LFLs "Operator Insights"-Modul (Reparierbarkeit, '
         'Kreislaufwirtschaft) und der Design-Copilot-Ansatz schaffen eine skalierbare digitale '
         'Wissensinfrastruktur — ein echter Whitespace.'),
    ]

    for i, (title, body) in enumerate(whitespace, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(f'{i}.  {title}')
        r.bold = True
        r.font.color.rgb = DARK_BLUE
        r.font.size = Pt(11)
        add_para(doc, body, space_after=4)

    doc.add_paragraph()

    # Closing statement
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    shade_cell_like_para = OxmlElement('w:pPr')
    run = p.add_run(
        '⟶  LoopForgeLab besetzt als einziger Anbieter den Whitespace der '
        '"economic integration" — die Verknüpfung von Engineering-Physik mit Wirtschaftslogik '
        'im Moment maximalen Einflusses: der frühen Designphase. Das ist nicht ein Feature, '
        'das Wettbewerber morgen kopieren können — es erfordert tiefes Domain-Know-how '
        'in Mechanical Engineering, Produktökonomie und KI-Engineering gleichzeitig.'
    )
    run.bold = True
    run.font.color.rgb = DARK_BLUE
    run.font.size = Pt(11)

    doc.add_paragraph()

    # ── Footer note ──────────────────────────
    add_para(doc, (
        '\n\nQuellenhinweis: Alle Marktdaten basieren auf der Competitive Intelligence Datenbank '
        'v1.7 (Stand: 13. März 2026). Qualitative Insights basieren auf dem LoopForgeLab Whitepaper '
        '(Februar 2026, 44 Experteninterviews) und dem LoopForgeLab Pitch Deck (März 2026). '
        'Externe Benchmarks: Loopio RfP Benchmark Report, Resilinc Supply Chain Disruption Study 2023, '
        'DriveWorks Engineering Change Studies, McKinsey Supply Chain Disruption Cost Analysis 2023, '
        'Center for Automotive Research (2026), CSRD-Regulierung der EU-Kommission.'
    ), size=7.5, color=RGBColor(0x80, 0x80, 0x90))

    # ── Save ─────────────────────────────────
    out_path = '/home/user/LFL_Competitor_Report/260313_LFL_Competitor_Report.docx'
    doc.save(out_path)
    print(f'✅  Report saved to: {out_path}')

if __name__ == '__main__':
    build_report()
